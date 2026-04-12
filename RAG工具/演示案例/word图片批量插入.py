# -*- coding: utf-8 -*-
"""
Word图片批量插入工具 - 现有文档版
功能：打开现有Word文档，将图片插入到指定表格的两个单元格中（每格一张）

作者：Claw
创建时间：2026-04-11
更新时间：2026-04-11
"""

import os
import re
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ==================== 配置区域（可调节）====================

CONFIG = {
    # 路径配置
    "source_folder": "./图片文件夹",      # 图片文件夹路径
    "template_file": "./模板文档.docx",    # 现有Word文档路径
    "output_file": "./输出文档.docx",      # 输出Word文件路径
    
    # 图片配置
    "image_height": 5.0,                   # 图片固定高度（cm）
    
    # 表格定位配置
    "target_row_position": -2,             # 目标行位置：-2表示倒数第二行
    # 自动找该行的最左和最右单元格，无需指定列索引
}

# ==========================================================


def get_sorted_folders(source_folder):
    """获取并排序文件夹列表"""
    if not os.path.exists(source_folder):
        raise FileNotFoundError(f"源文件夹不存在: {source_folder}")
    
    folders = []
    for item in os.listdir(source_folder):
        item_path = os.path.join(source_folder, item)
        if os.path.isdir(item_path):
            folders.append(item_path)
    
    def extract_number(folder_path):
        folder_name = os.path.basename(folder_path)
        match = re.match(r'^(\d+)', folder_name)
        return int(match.group(1)) if match else 999999
    
    folders.sort(key=extract_number)
    print(f"找到 {len(folders)} 个文件夹")
    return folders


def get_images_from_folder(folder_path):
    """从文件夹中获取图片文件（最多2张）"""
    images = []
    supported_extensions = ['.jpg', '.jpeg', '.png']
    
    for item in os.listdir(folder_path):
        item_path = os.path.join(folder_path, item)
        if os.path.isfile(item_path):
            ext = os.path.splitext(item)[1].lower()
            if ext in supported_extensions:
                images.append(item_path)
    
    images.sort()
    return images[:2]


def insert_image_to_cell(cell, image_path, height_cm):
    """在单元格中插入图片"""
    cell.paragraphs[0].clear()
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = cell.paragraphs[0].add_run()
    try:
        run.add_picture(image_path, height=Cm(height_cm))
        print(f"    已插入: {os.path.basename(image_path)}")
    except Exception as e:
        print(f"    插入失败: {os.path.basename(image_path)}, 错误: {e}")


def get_row_index(table, row_position):
    """根据位置参数获取行索引"""
    if row_position < 0:
        return len(table.rows) + row_position
    return row_position


def get_left_and_right_cells(table, row_idx):
    """
    获取某一行最左侧和最右侧的单元格
    
    Args:
        table: 表格对象
        row_idx: 行索引
    
    Returns:
        (left_cell, right_cell) 元组
    """
    row = table.rows[row_idx]
    cells = list(row.cells)
    
    # 去重（合并单元格会重复出现）
    unique_cells = []
    seen_tc = set()
    for cell in cells:
        tc = cell._tc
        if tc not in seen_tc:
            unique_cells.append(cell)
            seen_tc.add(tc)
    
    left_cell = unique_cells[0]
    right_cell = unique_cells[-1]
    
    return left_cell, right_cell


def main():
    """主函数"""
    print("=" * 50)
    print("Word图片批量插入工具 - 现有文档版")
    print("=" * 50)
    
    config = CONFIG
    source_folder = config["source_folder"]
    template_file = config["template_file"]
    output_file = config["output_file"]
    
    print(f"\n源文件夹: {source_folder}")
    print(f"模板文档: {template_file}")
    print(f"输出文件: {output_file}")
    print(f"目标行位置: {config['target_row_position']} (负数表示倒数)")
    print(f"图片高度: {config['image_height']} cm")
    
    # 检查文件
    if not os.path.exists(source_folder):
        print(f"\n错误: 源文件夹不存在: {source_folder}")
        return
    
    if not os.path.exists(template_file):
        print(f"\n错误: 模板文档不存在: {template_file}")
        return
    
    # 获取图片文件夹
    folders = get_sorted_folders(source_folder)
    if not folders:
        print("\n错误: 没有找到任何文件夹")
        return
    
    # 打开文档
    print(f"\n正在打开文档: {template_file}")
    doc = Document(template_file)
    
    tables = doc.tables
    print(f"文档中共有 {len(tables)} 个表格")
    
    if len(folders) > len(tables):
        print(f"警告: 图片组数量({len(folders)})大于表格数量({len(tables)})")
    
    # 处理
    print("\n开始插入图片...")
    processed, skipped = 0, 0
    
    for i, folder in enumerate(folders):
        if i >= len(tables):
            print(f"\n[{i+1}] 跳过: 表格数量不足")
            skipped += 1
            continue
        
        folder_name = os.path.basename(folder)
        print(f"\n[{i+1}/{len(folders)}] 处理: {folder_name} -> 表格{i+1}")
        
        images = get_images_from_folder(folder)
        
        if len(images) < 2:
            print(f"  跳过: 图片数量不足")
            skipped += 1
            continue
        
        table = tables[i]
        row_idx = get_row_index(table, config["target_row_position"])
        
        # 检查行索引
        if row_idx < 0 or row_idx >= len(table.rows):
            print(f"  跳过: 行索引 {row_idx} 超出范围")
            skipped += 1
            continue
        
        # 获取最左和最右单元格
        left_cell, right_cell = get_left_and_right_cells(table, row_idx)
        
        print(f"  目标行索引: {row_idx}, 找到 {len(images)} 张图片")
        
        # 插入左图
        try:
            insert_image_to_cell(left_cell, images[0], config["image_height"])
        except Exception as e:
            print(f"  左图插入失败: {e}")
        
        # 插入右图
        if len(images) >= 2:
            # 检查左右是否是同一个单元格（整行合并的情况）
            if left_cell._tc == right_cell._tc:
                print(f"  警告: 该行只有一个合并单元格，两张图片会叠在一起")
            try:
                insert_image_to_cell(right_cell, images[1], config["image_height"])
            except Exception as e:
                print(f"  右图插入失败: {e}")
        
        processed += 1
    
    # 保存
    print(f"\n正在保存: {output_file}")
    doc.save(output_file)
    
    print("\n" + "=" * 50)
    print(f"完成！成功: {processed}，跳过: {skipped}")
    print("=" * 50)


# ==================== 使用说明 ====================
"""
使用方法:
1. 修改配置:
   - template_file: 你的Word文档路径
   - target_row_position: 目标行（-2=倒数第二行）
   - target_col_left: 左图列索引（默认0，第1列）
   - target_col_right: 右图列索引（默认1，第2列）
   - image_height: 图片高度（cm）

2. 表格结构要求:
   - 每个表格的目标行有两列
   - 第1列放图1，第2列放图2

3. 运行:
   python word图片批量插入.py
"""
# ====================================================


if __name__ == "__main__":
    main()
